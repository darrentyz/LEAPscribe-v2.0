import streamlit as st
from core.auth import require_password
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(page_title="Summary & Download", page_icon="📦", layout="wide")
require_password()

st.title("📦 Step 6 — Summary & Download")

case_md = st.session_state.get("case_markdown")
cover_image = st.session_state.get("cover_image")
diagram_images = st.session_state.get("diagram_images", {})
selected_bps = st.session_state.get("selected_best_practices", [])

if not case_md:
    st.warning("Please complete the earlier steps first.")
    st.stop()

SECTION_ORDER = [
    "Title",
    "Executive Summary",
    "Problem / Need",
    "Implementation Approach",
    "Benefits & Impact",
    "Key Learning Points",
    "Point of Contact",
]

def parse_sections(md: str):
    sections = {"Title": []}
    current = "Title"
    for raw_line in md.splitlines():
        line = raw_line.strip()
        if line.startswith("# "):
            sections["Title"] = [line[2:].strip()]
            current = "Title"
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            normalized = None
            for std in SECTION_ORDER[1:]:
                if std.lower().split()[0] in heading.lower():
                    normalized = std
                    break
            heading_key = normalized or heading
            sections.setdefault(heading_key, [])
            current = heading_key
            continue
        sections.setdefault(current, []).append(raw_line)
    return sections

sections = parse_sections(case_md)
for sec in SECTION_ORDER:
    sections.setdefault(sec, [])

if "image_placement" not in st.session_state:
    mapping = {}
    if cover_image:
        mapping["__COVER__"] = "Title"
    for prompt in diagram_images.keys():
        mapping[prompt] = "Implementation Approach"
    st.session_state["image_placement"] = mapping

placement = st.session_state["image_placement"]

st.header("Preview")
if cover_image:
    st.image(cover_image, caption="Cover Image", use_column_width=True)
st.markdown(case_md)

st.divider()
st.header("Image Placement")

with st.expander("Cover Image Placement", expanded=True if cover_image else False):
    if cover_image:
        placement["__COVER__"] = st.selectbox(
            "Place cover image under section:",
            options=SECTION_ORDER,
            index=SECTION_ORDER.index(placement.get("__COVER__", "Title")),
            key="cover_select",
        )
    else:
        st.caption("No cover image generated in Step 4.")

if diagram_images:
    st.subheader("Diagram Placement")
    for prompt, img_bytes in diagram_images.items():
        cols = st.columns([3, 2])
        with cols[0]:
            st.image(img_bytes, caption=prompt, use_column_width=True)
        with cols[1]:
            current_sec = placement.get(prompt, "Implementation Approach")
            placement[prompt] = st.selectbox(
                f"Place diagram for:\n“{prompt[:60]}{'…' if len(prompt)>60 else ''}”",
                options=SECTION_ORDER,
                index=SECTION_ORDER.index(current_sec)
                if current_sec in SECTION_ORDER
                else 3,
                key=f"sel_{hash(prompt)}",
            )
else:
    st.caption("No diagrams generated in Step 4.")

st.session_state["image_placement"] = placement

st.divider()
st.header("Mapped Capability & Best Practice Statements")

if selected_bps:
    for bp in selected_bps:
        st.markdown(
            f"- **{bp['cap_id']} {bp['capability']}** – `{bp['metric_id']}`: {bp['statement']}"
        )
else:
    st.info(
        "No best practice statements selected yet. You can map them in **5️⃣ Map Capabilities**."
    )

def build_docx_from_sections(
    sections: dict,
    placement: dict,
    cover_img: bytes | None,
    diag_imgs: dict,
    selected_bps: list,
):
    doc = Document()

    def add_img(img_bytes: bytes):
        doc.add_picture(BytesIO(img_bytes), width=Inches(5.5))
        doc.add_paragraph()

    title_text = "Case Study"
    if sections.get("Title"):
        t = sections["Title"][0].strip("# ").strip()
        if t:
            title_text = t
    doc.add_heading(title_text, level=1)

    if cover_img and placement.get("__COVER__") == "Title":
        add_img(cover_img)

    rendered = set(["Title"])
    ordered_keys = [k for k in SECTION_ORDER if k != "Title"] + [
        k for k in sections.keys() if k not in SECTION_ORDER
    ]

    for key in ordered_keys:
        if key in rendered:
            continue
        if "visual" in key.lower() or "diagram" in key.lower():
            continue
        rendered.add(key)
        doc.add_heading(key, level=2)
        body_lines = sections.get(key, [])
        for raw in body_lines:
            line = raw.rstrip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
        if cover_img and placement.get("__COVER__") == key and key != "Title":
            add_img(cover_img)
        for prompt, img in (diag_imgs or {}).items():
            if placement.get(prompt) == key:
                add_img(img)

    if selected_bps:
        doc.add_page_break()
        doc.add_heading("Mapped Capability & Best Practice Statements", level=2)
        for bp in selected_bps:
            line = f"{bp['cap_id']} {bp['capability']} – {bp['metric_id']}: {bp['statement']}"
            doc.add_paragraph(line, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.divider()
if st.button("📥 Generate & Download DOCX"):
    doc_buf = build_docx_from_sections(
        sections=sections,
        placement=placement,
        cover_img=cover_image,
        diag_imgs=diagram_images,
        selected_bps=selected_bps,
    )
    st.success("✅ DOCX generated successfully!")
    st.download_button(
        label="⬇️ Download Case Study (.docx)",
        data=doc_buf,
        file_name="LEAPscribe_Case_Study.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
